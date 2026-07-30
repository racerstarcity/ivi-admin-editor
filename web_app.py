#!/usr/bin/env python3
"""Веб-интерфейс для генерации документов гостевого дома «Лидия»"""

import csv
import io
import os
import re
import sys
import zipfile
from urllib.request import urlopen, quote
from flask import Flask, render_template_string, request, send_file, jsonify

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SHEET_ID = "1cJ8HjD9cd-mTgtknUcjFs9cOgoBV28Jw"
SHEETS = [
    {"label": "2026", "gid": "351727279"},
    {"label": "2026 Наталья", "gid": "567545954"},
]

DOC_TYPES = [
    {"id": "receipt",  "label": "Квитанция"},
    {"id": "cert",     "label": "Справка Эконом"},
    {"id": "contract", "label": "Договор Лосинки"},
    {"id": "price",    "label": "Прайс-2500"},
    {"id": "egrip",    "label": "ЕГРИП"},
    {"id": "ip_inn",   "label": "ИП ИНН"},
]

def _template_path(name):
    if getattr(sys, 'frozen', False):
        base = sys._MEIPASS
        bundled = os.path.join(base, name)
        if os.path.exists(bundled):
            return bundled
    base = os.path.dirname(os.path.abspath(sys.argv[0]))
    local = os.path.join(base, name)
    if os.path.exists(local):
        return local
    desktop = os.path.join(os.path.expanduser("~"), "Desktop", name)
    if os.path.exists(desktop):
        return desktop
    return name

TEMPLATE_RECEIPT  = _template_path("Квитанция.docx")
TEMPLATE_CERT     = _template_path("Справка Эконом.docx")
TEMPLATE_CONTRACT = _template_path("Договор Лосинки.docx")
TEMPLATE_PRICE    = _template_path("Прайс-2500.docx")
TEMPLATE_EGRIP    = _template_path("ЕГРИП.pdf")
TEMPLATE_IP_INN   = _template_path("ИП ИНН.pdf")

app = Flask(__name__)


def sheet_url(gid):
    return f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv&gid={gid}"

_cache = {}

def load_data(gid=None):
    if gid is None:
        gid = SHEETS[0]["gid"]
    if gid in _cache:
        return _cache[gid]
    url = sheet_url(gid)
    resp = urlopen(url, timeout=15)
    text = resp.read().decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(text))
    rows = []
    for row in reader:
        row = {k.strip(): v.strip() for k, v in row.items() if k}
        if row.get("ФИО") and row.get("Цена") and row["Цена"] != "0":
            row["Цена"] = float(row["Цена"])
            row["Тариф"] = float(row["Тариф"]) if row.get("Тариф") else 0
            row["Сутки"] = int(row["Сутки"]) if row.get("Сутки") else 1
            rows.append(row)
    _cache[gid] = rows
    return rows


def replace_para_all(p, new_text):
    if not p.runs:
        p.add_run(new_text)
        return
    ref = p.runs[0]
    fn, fs, b = ref.font.name, ref.font.size, ref.bold
    for r in p.runs:
        r.text = ""
    p.runs[0].text = new_text
    if fn:
        p.runs[0].font.name = fn
    if fs:
        p.runs[0].font.size = fs
    if b is not None:
        p.runs[0].bold = b


def replace_in_para(p, old, new):
    """Заменяет подстроку в параграфе, сохраняя форматирование"""
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)


def fmt_bar(d):
    parts = d.split(".")
    return f"{parts[0]}/{parts[1]}/{parts[2][-2:]}" if len(parts) == 3 else d


def fmt_full(d):
    parts = d.split(".")
    return f"{parts[0]}/{parts[1]}/{parts[2]}" if len(parts) == 3 else d


def dative_surname(surname):
    if not surname:
        return surname
    if surname.endswith(("их", "ых")):
        return surname
    if surname.endswith(("ев", "ёв")):
        return surname + "у"
    if surname.endswith("ов"):
        return surname + "у"
    if surname.endswith("ин") or surname.endswith("ын"):
        return surname + "у"
    if surname.endswith("ский"):
        return surname[:-1] + "ому"
    if surname.endswith("цкий"):
        return surname[:-1] + "ому"
    if surname.endswith("ий"):
        return surname[:-1] + "ему"
    if surname.endswith("ой"):
        return surname[:-1] + "ому"
    if surname.endswith("ая"):
        return surname[:-1] + "ой"
    last = surname[-1]
    if last in "бвгджзклмнпрстфхчшщ":
        return surname + "у"
    return surname


def num_words(n):
    u = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    uf = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    te = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    ts = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hd = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    if n == 0:
        return "ноль"
    r = []
    th = n // 1000
    if th > 0:
        if th < 10:
            r.append(uf[th] + (" тысяча" if th == 1 else " тысячи" if th in [2, 3, 4] else " тысяч"))
        elif th < 20:
            r.append(te[th - 10] + " тысяч")
        else:
            r.append(ts[th // 10])
            uu = th % 10
            if uu > 0:
                r.append(uf[uu] + (" тысяча" if uu == 1 else " тысячи" if uu in [2, 3, 4] else " тысяч"))
            else:
                r.append("тысяч")
    n = n % 1000
    r.append(hd[n // 100])
    n %= 100
    if n >= 20:
        r.append(ts[n // 10])
        n %= 10
        if n > 0:
            r.append(u[n])
    elif n >= 10:
        r.append(te[n - 10])
    elif n > 0:
        r.append(u[n])
    return " ".join(filter(None, r))


def gen_receipt(data):
    fio = data["ФИО"]
    account = data.get("Номер счета", "")
    num = account[-3:] if account else ""
    arrival = data["Дата приезда"]
    departure = data["Дата отьезда"]
    days = int(data["Сутки"])
    tariff = int(data["Тариф"])
    total = int(data["Цена"])
    date_str = fmt_full(data.get("Дата счета", arrival))
    arr_d = fmt_bar(arrival)
    dep_d = fmt_bar(departure)

    doc = Document(TEMPLATE_RECEIPT)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Дождиков Антон"):
            replace_para_all(p, fio)
        elif t.startswith("Приезд"):
            replace_para_all(p, f"Приезд\t\t{arr_d}")
        elif t.startswith("Отъезд"):
            replace_para_all(p, f"Отъезд\t\t{dep_d}")
        elif t.startswith("Сутки") or "Сутки" in t:
            replace_para_all(p, f"Сутки      \t      {days}\t\tТариф\t\t\t{tariff:.2f}\t")
        elif t.startswith("С Ч Е Т") or t.startswith("СЧЕТ"):
            replace_para_all(p, f"С Ч Е Т  {account}")
        elif "Всего:" in t and "оплачено" not in t and "к оплате" not in t:
            replace_para_all(p, f"\tВсего:\t\t{total:.2f} ₽")
        elif "Всего оплачено" in t:
            replace_para_all(p, f"\tВсего оплачено:   {total:.2f} ₽")
        elif "Всего к оплате" in t:
            replace_para_all(p, f"\tВсего к оплате:   {total:.2f} ₽")
        elif "Номер" in t and len(t) < 40:
            new = re.sub(r"\d+$", num, p.text) if num else p.text
            replace_para_all(p, new)
        elif "/" in t and len(t.strip()) <= 12 and "Приезд" not in t and "Отъезд" not in t and "Сутки" not in t:
            replace_para_all(p, date_str)

    for table in doc.tables:
        for row in table.rows:
            full = " | ".join(cell.text for cell in row.cells)
            if "Проживание" in full and "Вид начисления" not in full:
                cells = row.cells
                dr = f"{arr_d}-{dep_d}"
                ts = f"{total:.2f}"
                new_texts = ["Проживание, номер одноместный, однокомнатный «Стандарт»", dr, ts, ts]
                for ci, txt in enumerate(new_texts):
                    p = cells[ci].paragraphs[0]
                    if not p.runs:
                        p.add_run(txt)
                    else:
                        p.runs[0].text = txt
                        for rn in p.runs[1:]:
                            rn.text = ""

    fio_short = fio.replace(" ", "_")
    date_short = arrival.replace(".", "")
    fn = f"Квитанция_{fio_short}_{date_short}.docx"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return fn, buf


def gen_cert(data):
    fio = data["ФИО"]
    arrival = data["Дата приезда"]
    departure = data["Дата отьезда"]
    tariff = int(data["Тариф"])
    days = int(data["Сутки"])

    parts = fio.split()
    surname = parts[0] if len(parts) > 0 else ""
    name = parts[1] if len(parts) > 1 else ""
    patr = parts[2] if len(parts) > 2 else ""
    init = (name[0] + "." if name else "") + (patr[0] + "." if patr else "")
    surname_dat = dative_surname(surname)
    fio_d = f"{surname_dat} {init}" if init else surname_dat
    price_w = num_words(tariff)
    text = (
        f"Дана {fio_d} в том, что он с {arrival} по {departure} "
        f"проживал в однокомнатном одноместном номере «Стандарт» "
        f"гостевого дома «Лидия». "
        f"Стоимость проживания за "
        f"{'сутки' if days == 1 else 'суток'} в номере составляет "
        f"{tariff} ({price_w}) рублей."
    )
    doc = Document(TEMPLATE_CERT)
    for p in doc.paragraphs:
        if p.text.strip().startswith("Дана") and "Дождикову" in p.text:
            replace_para_all(p, text)

    fio_short = fio.replace(" ", "_")
    date_short = arrival.replace(".", "")
    fn = f"Справка_Эконом_{fio_short}_{date_short}.docx"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return fn, buf


def _copy_template(path, filename):
    """Скопировать шаблон как есть"""
    buf = io.BytesIO()
    with open(path, "rb") as f:
        buf.write(f.read())
    buf.seek(0)
    return filename, buf

gen_contract = lambda data: _copy_template(TEMPLATE_CONTRACT, "Договор Лосинки.docx")
gen_price     = lambda data: _copy_template(TEMPLATE_PRICE,     "Прайс-2500.docx")
gen_egrip     = lambda data: _copy_template(TEMPLATE_EGRIP,     "ЕГРИП.pdf")
gen_ip_inn    = lambda data: _copy_template(TEMPLATE_IP_INN,    "ИП ИНН.pdf")


def generate_docs(data, doc_ids):
    """Генерирует список (filename, buf) для запрошенных типов документов"""
    results = []
    for did in doc_ids:
        try:
            if did == "receipt":
                results.append(gen_receipt(data))
            elif did == "cert":
                results.append(gen_cert(data))
            elif did == "contract":
                results.append(gen_contract(data))
            elif did == "price":
                results.append(gen_price(data))
            elif did == "egrip":
                results.append(gen_egrip(data))
            elif did == "ip_inn":
                results.append(gen_ip_inn(data))
        except Exception as e:
            print(f"Ошибка генерации {did}: {e}")
    return results


HTML = """\
<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Генератор документов</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f0f2f5; color: #1a1a2e; padding: 20px; }
  .container { max-width: 1200px; margin: 0 auto; }
  h1 { font-size: 24px; margin-bottom: 5px; }
  .subtitle { color: #666; margin-bottom: 20px; font-size: 14px; }
  .search-box { margin-bottom: 20px; display: flex; gap: 10px; align-items: center; flex-wrap: wrap; }
  .search-box input { flex: 1; min-width: 200px; padding: 10px 14px; border: 1px solid #d0d5dd; border-radius: 8px; font-size: 15px; }
  .search-box input:focus { outline: none; border-color: #4f46e5; box-shadow: 0 0 0 3px rgba(79,70,229,.1); }
  .badge { background: #e0e7ff; color: #3730a3; padding: 4px 10px; border-radius: 20px; font-size: 13px; }
  .loading { display: none; margin-left: 10px; color: #666; }
  .loading.active { display: inline; }
  .tabs { display: flex; gap: 8px; margin-bottom: 20px; flex-wrap: wrap; }
  .tab { padding: 8px 20px; border-radius: 8px; border: 1px solid #d0d5dd; background: #fff; cursor: pointer; font-size: 14px; font-weight: 500; transition: all .15s; }
  .tab:hover { border-color: #4f46e5; }
  .tab.active { background: #4f46e5; color: #fff; border-color: #4f46e5; }
  .table-wrap { overflow-x: auto; background: #fff; border-radius: 12px; box-shadow: 0 1px 3px rgba(0,0,0,.08); }
  table { width: 100%; border-collapse: collapse; font-size: 14px; }
  th { background: #f8fafc; text-align: left; padding: 12px 14px; font-weight: 600; color: #475569; border-bottom: 1px solid #e2e8f0; white-space: nowrap; }
  td { padding: 10px 14px; border-bottom: 1px solid #f1f5f9; }
  tr:hover td { background: #f8fafc; }
  .fio { font-weight: 500; }
  .dates { color: #64748b; font-size: 13px; }
  .sum { font-weight: 600; color: #059669; }
  .btn { display: inline-flex; align-items: center; gap: 6px; padding: 7px 16px; border: none; border-radius: 8px; font-size: 13px; font-weight: 500; cursor: pointer; text-decoration: none; transition: all .15s; }
  .btn-primary { background: #4f46e5; color: #fff; }
  .btn-primary:hover { background: #4338ca; }
  .btn-outline { background: transparent; color: #4f46e5; border: 1px solid #4f46e5; }
  .btn-outline:hover { background: #eef2ff; }
  .btn-success { background: #059669; color: #fff; }
  .btn-success:hover { background: #047857; }
  .btn:disabled { opacity: .5; cursor: default; }
  .actions { display: flex; gap: 6px; }
  .check-cell { width: 40px; text-align: center; }
  .check-cell input { width: 18px; height: 18px; cursor: pointer; }
  .toolbar { margin-bottom: 16px; display: flex; gap: 10px; align-items: center; flex-wrap: wrap; }
  .doc-checkboxes { display: flex; gap: 12px; flex-wrap: wrap; margin-bottom: 16px; padding: 12px 16px; background: #fff; border-radius: 10px; box-shadow: 0 1px 2px rgba(0,0,0,.06); }
  .doc-checkboxes label { display: inline-flex; align-items: center; gap: 6px; cursor: pointer; font-size: 14px; }
  .doc-checkboxes input { width: 17px; height: 17px; cursor: pointer; }
  .toast { position: fixed; bottom: 30px; left: 50%; transform: translateX(-50%); background: #1a1a2e; color: #fff; padding: 12px 24px; border-radius: 10px; font-size: 14px; opacity: 0; transition: opacity .3s; pointer-events: none; z-index: 100; }
  .toast.show { opacity: 1; }
  .empty { text-align: center; padding: 40px; color: #94a3b8; }
  @media (max-width: 768px) {
    .search-box { flex-direction: column; }
    .search-box input { width: 100%; }
    .toolbar { flex-direction: column; }
    .btn { width: 100%; justify-content: center; }
  }
</style>
</head>
<body>
<div class="container">
  <h1>Генератор документов</h1>
  <p class="subtitle">Гостевой дом «Лидия»</p>

  <div class="tabs" id="tabs"></div>

  <div class="doc-checkboxes" id="docTypes">
    <label><input type="checkbox" value="receipt" checked onchange="updateGenBtn()"> Квитанция</label>
    <label><input type="checkbox" value="cert" checked onchange="updateGenBtn()"> Справка Эконом</label>
    <label><input type="checkbox" value="contract" onchange="updateGenBtn()"> Договор Лосинки</label>
    <label><input type="checkbox" value="price" onchange="updateGenBtn()"> Прайс-2500</label>
    <label><input type="checkbox" value="egrip" onchange="updateGenBtn()"> ЕГРИП</label>
    <label><input type="checkbox" value="ip_inn" onchange="updateGenBtn()"> ИП ИНН</label>
  </div>

  <div class="search-box">
    <input type="text" id="search" placeholder="Поиск по ФИО..." oninput="filterTable()">
    <span class="badge" id="count">0</span>
    <button class="btn btn-outline" onclick="refreshData()" style="font-size:13px;padding:5px 12px">⟳ Обновить</button>
    <span class="loading" id="loading">Загрузка...</span>
  </div>

  <div class="toolbar">
    <button class="btn btn-primary" onclick="genSelected()" id="genBtn" disabled>Скачать выбранные (ZIP)</button>
    <button class="btn btn-success" onclick="genAll()">Скачать все (ZIP)</button>
    <label style="display:inline-flex;align-items:center;gap:6px;cursor:pointer;font-size:13px">
      <input type="checkbox" onchange="toggleAll(this)"> Выбрать все
    </label>
  </div>

  <div class="table-wrap">
    <table>
      <thead>
        <tr>
          <th class="check-cell"></th>
          <th>ФИО</th>
          <th>Заезд</th>
          <th>Выезд</th>
          <th>Суток</th>
          <th>Тариф</th>
          <th>Сумма</th>
          <th>Номер счёта</th>
          <th>Действия</th>
        </tr>
      </thead>
      <tbody id="tbody"></tbody>
    </table>
  </div>
</div>
<div class="toast" id="toast"></div>

<script>
const SHEETS = [
  {label: "2026", gid: "351727279"},
  {label: "2026 Наталья", gid: "567545954"},
];

let currentGid = SHEETS[0].gid;
let rows = [];

function initTabs() {
  const container = document.getElementById('tabs');
  container.innerHTML = '';
  for (const s of SHEETS) {
    const div = document.createElement('div');
    div.className = 'tab' + (s.gid === currentGid ? ' active' : '');
    div.textContent = s.label;
    div.dataset.gid = s.gid;
    div.onclick = () => switchTab(s.gid);
    container.appendChild(div);
  }
}

function getSelectedDocs() {
  return Array.from(document.querySelectorAll('#docTypes input:checked')).map(c => c.value);
}

async function switchTab(gid) {
  currentGid = gid;
  document.querySelectorAll('.tab').forEach(t => t.classList.toggle('active', t.dataset.gid === gid));
  await loadData();
}

async function loadData() {
  const loading = document.getElementById('loading');
  loading.classList.add('active');
  try {
    const r = await fetch('/api/rows?gid=' + currentGid, {signal: AbortSignal.timeout(20000)});
    if (!r.ok) { const e = await r.json(); showToast('Ошибка: ' + (e.error || r.status)); return; }
    rows = await r.json();
  } catch(e) { showToast('Ошибка загрузки: ' + (e.message || 'таймаут')); }
  loading.classList.remove('active');
  renderTable();
}

function renderTable() {
  const q = document.getElementById('search').value.toLowerCase();
  const filtered = rows.filter(r => r.fio.toLowerCase().includes(q));
  const tbody = document.getElementById('tbody');
  tbody.innerHTML = '';
  document.getElementById('count').textContent = filtered.length + ' записей';
  if (filtered.length === 0) {
    tbody.innerHTML = '<tr><td colspan="9" class="empty">Ничего не найдено</td></tr>';
    return;
  }
  for (const r of filtered) {
    const tr = document.createElement('tr');
    tr.dataset.idx = r.idx;
    tr.innerHTML = `
      <td class="check-cell"><input type="checkbox" class="row-check" data-idx="${r.idx}" onchange="updateGenBtn()"></td>
      <td class="fio">${r.fio}</td>
      <td class="dates">${r.arrival}</td>
      <td class="dates">${r.departure}</td>
      <td>${r.days}</td>
      <td>${r.tariff}₽</td>
      <td class="sum">${r.total}₽</td>
      <td style="font-size:12px;color:#64748b">${r.account}</td>
      <td class="actions">
        <button class="btn btn-primary" onclick="downloadOne(${r.idx})">Скачать</button>
      </td>
    `;
    tbody.appendChild(tr);
  }
}

function filterTable() { renderTable(); }

function toggleAll(cb) {
  document.querySelectorAll('.row-check').forEach(c => c.checked = cb.checked);
  updateGenBtn();
}

async function refreshData() {
  const loading = document.getElementById('loading');
  loading.classList.add('active');
  try {
    const r = await fetch('/api/clear_cache?gid=' + currentGid);
    if (r.ok) showToast('Данные обновлены');
  } catch(e) {}
  await loadData();
}

function updateGenBtn() {
  const checked = document.querySelectorAll('.row-check:checked').length;
  document.getElementById('genBtn').disabled = checked === 0;
}

function showToast(msg) {
  const t = document.getElementById('toast');
  t.textContent = msg;
  t.classList.add('show');
  setTimeout(() => t.classList.remove('show'), 3000);
}

function downloadOne(idx) {
  const url = '/download/' + idx + '?gid=' + currentGid + '&docs=' + getSelectedDocs().join(',');
  window.open(url, '_blank');
  showToast('Скачивание...');
}

function genSelected() {
  const checked = document.querySelectorAll('.row-check:checked');
  if (checked.length === 0) return;
  const idxs = Array.from(checked).map(c => c.dataset.idx);
  const url = '/download_batch?ids=' + idxs.join(',') + '&gid=' + currentGid + '&docs=' + getSelectedDocs().join(',');
  window.open(url, '_blank');
  showToast('Скачивание...');
}

function genAll() {
  const url = '/download_all?gid=' + currentGid + '&docs=' + getSelectedDocs().join(',');
  window.open(url, '_blank');
  showToast('Скачивание...');
}

initTabs();
loadData();
</script>
</body>
</html>
"""


@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/api/rows")
def api_rows():
    gid = request.args.get("gid", SHEETS[0]["gid"])
    try:
        data = load_data(gid)
    except Exception as e:
        return jsonify({"error": str(e)}), 502
    return jsonify([
        {"idx": i, "fio": r["ФИО"], "arrival": r["Дата приезда"], "departure": r["Дата отьезда"],
         "days": r["Сутки"], "tariff": int(r["Тариф"]), "total": int(r["Цена"]),
         "account": r.get("Номер счета", "")}
        for i, r in enumerate(data)
    ])


@app.route("/api/clear_cache")
def clear_cache():
    gid = request.args.get("gid", "")
    if gid in _cache:
        del _cache[gid]
    return jsonify({"ok": True})


@app.route("/download/<int:idx>")
def download_one(idx):
    gid = request.args.get("gid", SHEETS[0]["gid"])
    docs_str = request.args.get("docs", "receipt,cert")
    doc_ids = [d.strip() for d in docs_str.split(",") if d.strip()]
    data = load_data(gid)
    if idx < 0 or idx >= len(data):
        return "Неверный индекс", 404
    files = generate_docs(data[idx], doc_ids)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fn, buf in files:
            zf.writestr(fn, buf.read())
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype="application/zip", as_attachment=True, download_name=f"документы_{idx}.zip")


SHARED_DOC_IDS = {"contract", "price", "egrip", "ip_inn"}

def _add_shared_docs(zf, doc_ids):
    """Добавляет общие документы (договор, прайс) один раз"""
    shared = [d for d in doc_ids if d in SHARED_DOC_IDS]
    for fn, buf in generate_docs({}, shared):
        zf.writestr(fn, buf.read())

@app.route("/download_batch")
def download_batch():
    ids = request.args.get("ids", "")
    gid = request.args.get("gid", SHEETS[0]["gid"])
    docs_str = request.args.get("docs", "receipt,cert")
    doc_ids = [d.strip() for d in docs_str.split(",") if d.strip()]
    per_guest = [d for d in doc_ids if d not in SHARED_DOC_IDS]
    data = load_data(gid)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for s in ids.split(","):
            if not s.strip():
                continue
            idx = int(s.strip())
            if 0 <= idx < len(data):
                prefix = f"{idx:04d}_"
                for fn, buf in generate_docs(data[idx], per_guest):
                    zf.writestr(prefix + fn, buf.read())
        _add_shared_docs(zf, doc_ids)
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype="application/zip", as_attachment=True, download_name="выбранные_документы.zip")


@app.route("/download_all")
def download_all():
    gid = request.args.get("gid", SHEETS[0]["gid"])
    docs_str = request.args.get("docs", "receipt,cert")
    doc_ids = [d.strip() for d in docs_str.split(",") if d.strip()]
    per_guest = [d for d in doc_ids if d not in SHARED_DOC_IDS]
    data = load_data(gid)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, row in enumerate(data):
            try:
                prefix = f"{i+1:04d}_"
                for fn, buf in generate_docs(row, per_guest):
                    zf.writestr(prefix + fn, buf.read())
            except Exception as e:
                print(f"Ошибка {i}: {e}")
        _add_shared_docs(zf, doc_ids)
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype="application/zip", as_attachment=True, download_name="все_документы.zip")


if __name__ == "__main__":
    import webbrowser
    port = 5000
    webbrowser.open(f"http://localhost:{port}")
    app.run(host="127.0.0.1", port=port, debug=False)
