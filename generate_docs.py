#!/usr/bin/env python3
"""
Генератор квитанций и справок на основе шаблонов .docx
"""

import argparse
import csv
import io
import os
import copy
import re
from urllib.request import urlopen

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SHEET_URL = "https://docs.google.com/spreadsheets/d/1cJ8HjD9cd-mTgtknUcjFs9cOgoBV28Jw/export?format=csv&gid=351727279"
TEMPLATE_RECEIPT = r"C:\Users\user\Desktop\Квитанция.docx"
TEMPLATE_CERT = r"C:\Users\user\Desktop\Справка Эконом.docx"

DIRECTOR = "В. В. Кравченко"


def load_data(url=None):
    if url is None:
        url = SHEET_URL
    resp = urlopen(url)
    text = resp.read().decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(text))
    rows = []
    for row in reader:
        row = {k.strip(): v.strip() for k, v in row.items() if k}
        if row.get("ФИО") and row.get("Цена") and row["Цена"] != "0":
            row["Цена"] = float(row["Цена"])
            row["Тариф"] = float(row["Тариф"]) if row.get("Тариф") else 0
            row["Сутки"] = int(row["Сутки"]) if row.get("Сутки") else 1
            rows.append(row)
    return rows


def find_record(rows, fio=None, date=None, account=None):
    matches = rows
    if fio:
        matches = [r for r in matches if fio.lower() in r["ФИО"].lower()]
    if date:
        matches = [r for r in matches if date in r["Дата приезда"]]
    if account:
        matches = [r for r in matches if account in r.get("Номер счета", "")]
    return matches


def replace_para_text(para, new_text, keep_first_run_style=True):
    """Заменяет текст в параграфе, сохраняя стиль первого run"""
    if keep_first_run_style and para.runs:
        ref = para.runs[0]
        font_name = ref.font.name
        font_size = ref.font.size
        bold = ref.bold
        italic = ref.italic
        underline = ref.underline
    else:
        font_name = None
        font_size = None
        bold = None
        italic = None
        underline = None

    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = new_text
        if font_name:
            para.runs[0].font.name = font_name
        if font_size:
            para.runs[0].font.size = font_size
        if bold is not None:
            para.runs[0].bold = bold
        if italic is not None:
            para.runs[0].italic = italic
        if underline is not None:
            para.runs[0].underline = underline
    else:
        para.runs[0].text = new_text


def replace_para_text_all(para, new_text):
    """Очищает все run'ы и добавляет один с новым текстом, сохраняя шрифт первого"""
    if not para.runs:
        para.add_run(new_text)
        return
    ref = para.runs[0]
    fn = ref.font.name
    fs = ref.font.size
    b = ref.bold
    for run in para.runs:
        run.text = ""
    para.runs[0].text = new_text
    if fn:
        para.runs[0].font.name = fn
    if fs:
        para.runs[0].font.size = fs
    if b is not None:
        para.runs[0].bold = b


def fmt_date_bar(d):
    """DD.MM.YYYY -> DD/MM/YY"""
    parts = d.split(".")
    if len(parts) == 3:
        return f"{parts[0]}/{parts[1]}/{parts[2][-2:]}"
    return d


def fmt_date_slash_full(d):
    """DD.MM.YYYY -> DD/MM/YYYY"""
    parts = d.split(".")
    if len(parts) == 3:
        return f"{parts[0]}/{parts[1]}/{parts[2]}"
    return d


def dative_surname(surname):
    if not surname:
        return surname
    if surname.endswith(("их", "ых")):
        return surname
    if surname.endswith(("ев", "ёв")):
        return surname + "у"
    if surname.endswith("ов"):
        return surname + "у"
    if surname.endswith("ин") or surname.endswith("ын"):
        return surname + "у"
    if surname.endswith("ский"):
        return surname[:-1] + "ому"
    if surname.endswith("цкий"):
        return surname[:-1] + "ому"
    if surname.endswith("ий"):
        return surname[:-1] + "ему"
    if surname.endswith("ой"):
        return surname[:-1] + "ому"
    if surname.endswith("ая"):
        return surname[:-1] + "ой"
    last = surname[-1]
    if last in "бвгджзклмнпрстфхчшщ":
        return surname + "у"
    return surname


def num_to_words(n):
    units = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    units_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    if n == 0:
        return "ноль"
    result = []
    th = n // 1000
    if th > 0:
        if th < 10:
            result.append(units_f[th] + (" тысяча" if th == 1 else " тысячи" if th in [2, 3, 4] else " тысяч"))
        elif th < 20:
            result.append(teens[th - 10] + " тысяч")
        else:
            result.append(tens[th // 10])
            u = th % 10
            if u > 0:
                result.append(units_f[u] + (" тысяча" if u == 1 else " тысячи" if u in [2, 3, 4] else " тысяч"))
            else:
                result.append("тысяч")
    n = n % 1000
    result.append(hundreds[n // 100])
    n = n % 100
    if n >= 20:
        result.append(tens[n // 10])
        n = n % 10
        if n > 0:
            result.append(units[n])
    elif n >= 10:
        result.append(teens[n - 10])
        n = 0
    elif n > 0:
        result.append(units[n])
    return " ".join(filter(None, result))


def generate_receipt(data, output_path):
    doc = Document(TEMPLATE_RECEIPT)
    paras = doc.paragraphs

    fio = data["ФИО"]
    account = data.get("Номер счета", "")
    num = account[-3:] if account else ""
    arrival = data["Дата приезда"]
    departure = data["Дата отьезда"]
    days = int(data["Сутки"])
    tariff = int(data["Тариф"])
    total = int(data["Цена"])
    date_str = fmt_date_slash_full(data.get("Дата счета", arrival))
    arrival_d = fmt_date_bar(arrival)
    departure_d = fmt_date_bar(departure)

    # Собираем всё в список параграфов, заменяем по шаблону
    # Ищем параграфы по их начальному тексту

    for p in paras:
        text = p.text

        if text.strip() == "Г-н":
            pass  # не меняем

        elif text.strip().startswith("Дождиков Антон Сергеевич"):
            replace_para_text_all(p, fio)

        elif "/" in text and len(text.strip()) <= 12 and not any(x in text for x in ["Приезд", "Отъезд", "Сутки", "Номер"]):
            # Дата счета (строка с датой в центре)
            replace_para_text_all(p, date_str)

        elif "\tНомер" in text or text.strip().startswith("Номер") or "Номер" in text:
            # Строка с номером
            new = re.sub(r"\d+$", num, text) if num else text
            replace_para_text_all(p, new)

        elif text.strip().startswith("Приезд"):
            replace_para_text_all(p, f"Приезд\t\t{arrival_d}")

        elif text.strip().startswith("Отъезд"):
            replace_para_text_all(p, f"Отъезд\t\t{departure_d}")

        elif text.strip().startswith("Сутки") or "Сутки" in text:
            # Замена суток и тарифа
            tariff_str = f"{tariff:.2f}"
            new = f"Сутки      \t      {days}\t\tТариф\t\t\t{tariff_str}\t"
            replace_para_text_all(p, new)

        elif text.strip().startswith("С Ч Е Т") or text.strip().startswith("СЧЕТ"):
            replace_para_text_all(p, f"С Ч Е Т  {account}")

        elif "Всего:" in text and "оплачено" not in text and "к оплате" not in text:
            total_str = f"{total:.2f}"
            replace_para_text_all(p, f"\tВсего:\t\t{total_str} ₽")

        elif "Всего оплачено" in text:
            total_str = f"{total:.2f}"
            replace_para_text_all(p, f"\tВсего оплачено:   {total_str} ₽")

        elif "Всего к оплате" in text:
            total_str = f"{total:.2f}"
            replace_para_text_all(p, f"\tВсего к оплате:   {total_str} ₽")

        elif "Дежурный администратор" in text:
            pass  # не меняем

    # Таблица
    for table in doc.tables:
        for row in table.rows:
            full = " | ".join(cell.text for cell in row.cells)
            if "Проживание" in full and "Вид начисления" not in full:
                cells = row.cells
                if len(cells) >= 4:
                    dr = f"{arrival_d}-{departure_d}"
                    ts = f"{total:.2f}"
                    new_texts = [
                        "Проживание, номер одноместный, однокомнатный «Стандарт»",
                        dr, ts, ts
                    ]
                    for ci, txt in enumerate(new_texts):
                        p = cells[ci].paragraphs[0]
                        if not p.runs:
                            p.add_run(txt)
                        else:
                            p.runs[0].text = txt
                            for r in p.runs[1:]:
                                r.text = ""

    doc.save(output_path)


def generate_certificate(data, output_path):
    doc = Document(TEMPLATE_CERT)
    paras = doc.paragraphs

    fio = data["ФИО"]
    arrival = data["Дата приезда"]
    departure = data["Дата отьезда"]
    tariff = int(data["Тариф"])
    days = int(data["Сутки"])

    parts = fio.split()
    surname = parts[0] if len(parts) > 0 else ""
    name = parts[1] if len(parts) > 1 else ""
    patronymic = parts[2] if len(parts) > 2 else ""
    initials = ""
    if name:
        initials += name[0] + "."
    if patronymic:
        initials += patronymic[0] + "."
    surname_dat = dative_surname(surname)
    fio_dat = f"{surname_dat} {initials}" if initials else surname_dat

    price_words = num_to_words(tariff)
    price_text = f"{tariff} ({price_words})"

    text_main = (
        f"Дана {fio_dat} в том, что он с {arrival} по {departure} "
        f"проживал в однокомнатном одноместном номере «Стандарт» "
        f"гостевого дома «Лидия». "
        f"Стоимость проживания за "
        f"{'сутки' if days == 1 else 'суток'} в номере составляет "
        f"{price_text} рублей."
    )

    for p in paras:
        text = p.text.strip()

        if text.startswith("Дана") and "Дождикову" in text:
            replace_para_text_all(p, text_main)

        elif text.startswith("ИП Кравченко") or text.startswith("ИП"):
            pass  # шапку не меняем, она статична

        elif text == "С П Р А В К А":
            pass

        elif text.startswith("На момент заселения"):
            pass

        elif text.startswith("В гостевом доме"):
            pass

        elif text.startswith("Дополнительные услуги"):
            pass

        elif "Генеральный директор" in text:
            pass

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Генератор квитанций и справок")
    parser.add_argument("--fio", help="ФИО для поиска")
    parser.add_argument("--date", help="Дата заезда для поиска")
    parser.add_argument("--account", help="Номер счета для поиска")
    parser.add_argument("--list", action="store_true", help="Список записей")
    parser.add_argument("--out", default=".", help="Папка для сохранения")
    parser.add_argument("--url", default=SHEET_URL, help="URL Google Sheets CSV")

    args = parser.parse_args()

    rows = load_data(args.url)

    if args.list:
        for i, r in enumerate(rows, 1):
            acct = r.get("Номер счета", "") or ""
            print(f"{i:3d}. {r['ФИО'][:30]:30s} | {r['Дата приезда']} - {r['Дата отьезда']} | {int(r['Цена'])}₽ | {acct}")
        return

    if not (args.fio or args.date or args.account):
        print("Укажите --fio, --date, --account или --list")
        return

    matches = find_record(rows, fio=args.fio, date=args.date, account=args.account)

    if not matches:
        print("Записей не найдено")
        return

    if len(matches) > 1:
        print(f"Найдено {len(matches)} записей. Уточните поиск:")
        for i, r in enumerate(matches, 1):
            acct = r.get("Номер счета", "") or ""
            print(f"{i:3d}. {r['ФИО'][:30]:30s} | {r['Дата приезда']} - {r['Дата отьезда']} | Суток: {r['Сутки']} | {int(r['Цена'])}₽ | {acct}")
        return

    data = matches[0]
    fio_short = data["ФИО"].replace(" ", "_")
    date_short = data["Дата приезда"].replace(".", "")

    path_r = os.path.join(args.out, f"Квитанция_{fio_short}_{date_short}.docx")
    generate_receipt(data, path_r)
    print(f"Квитанция: {path_r}")

    path_c = os.path.join(args.out, f"Справка_Эконом_{fio_short}_{date_short}.docx")
    generate_certificate(data, path_c)
    print(f"Справка:   {path_c}")


if __name__ == "__main__":
    main()
